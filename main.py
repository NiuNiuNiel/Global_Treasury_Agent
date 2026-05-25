# Global Treasury Agent - Frontend Dashboard
# main.py

import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, filedialog
import threading
import os
import re
from datetime import datetime

# Try to import backend modules; fall back gracefully if unavailable
try:
    from Data_Retrieval import Database_Connector
    _DB_AVAILABLE = True
except ImportError:
    _DB_AVAILABLE = False

try:
    from Agent import Agent
    _AGENT_AVAILABLE = True
except ImportError:
    _AGENT_AVAILABLE = False

# ── Appearance ────────────────────────────────────────────────────────────────
ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("blue")

CONFIDENCE_THRESHOLD = 0.95  # >= 95% => auto-validated (0-1 scale, matching DB storage)

C = {
    "bg":         "#0e0e1a",
    "panel":      "#14142a",
    "card":       "#1a1a30",
    "row_even":   "#181828",
    "row_odd":    "#1d1d30",
    "row_sel":    "#252545",
    "border":     "#2a2a45",
    "accent":     "#4f8ef7",
    "accent_dim": "#3a6ec4",
    "text":       "#e0e0f0",
    "text_dim":   "#8888aa",
    "text_muted": "#50506a",
}

STATUS_META = {
    "pending":            {"label": "●  Pending",          "fg": "#f5a623"},
    "processing":         {"label": "⟳  Processing...",    "fg": "#4f8ef7"},
    "auto_validated":     {"label": "✓  Auto-Validated",   "fg": "#2ecc71"},
    "needs_manual":       {"label": "⚠  Needs Review",     "fg": "#e67e22"},
    "manually_validated": {"label": "✓  Approved",         "fg": "#2ecc71"},
    "error":              {"label": "✗  Error",            "fg": "#e74c3c"},
}

COLUMNS = [
    ("select",      "",              44,   "center"),
    ("invoice_id",  "Invoice ID",   140,   "w"),
    ("date",        "Date",         105,   "center"),
    ("amount",      "Amount",       110,   "e"),
    ("currency",    "CCY",           58,   "center"),
    ("description", "Description",  235,   "w"),
    ("confidence",  "Confidence",   100,   "center"),
    ("status",      "Status",       165,   "center"),
    ("actions",     "Actions",      145,   "center"),
]

# Supported upload extensions -> (file_type stored in DB, requires_ocr)
UPLOAD_EXT_MAP = {
    "pdf":  ("pdf",  False),
    "docx": ("docx", False),
    "jpg":  ("jpg",  True),
    "jpeg": ("jpg",  True),
    "png":  ("png",  True),
}

MOCK_INVOICES = [
    {
        "invoice_id": "INV-0001", "_db_id": None, "_confidence": None, "_matched_ids": [],
        "date_of_purchase": "2026-05-10", "amount": 15000.00, "currency": "USD",
        "description": "Software licensing Q2 - Acme Corp",
        "validation_status": "pending",
    },
    {
        "invoice_id": "INV-0002", "_db_id": None, "_confidence": None, "_matched_ids": [],
        "date_of_purchase": "2026-05-12", "amount": 4200.50, "currency": "EUR",
        "description": "Cloud infrastructure - Globex Solutions",
        "validation_status": "pending",
    },
    {
        "invoice_id": "INV-0003", "_db_id": None, "_confidence": 0.82, "_matched_ids": [],
        "date_of_purchase": "2026-05-14", "amount": 87500.00, "currency": "MYR",
        "description": "Consulting retainer - TechServe Sdn Bhd",
        "validation_status": "needs_manual",
    },
    {
        "invoice_id": "INV-0004", "_db_id": None, "_confidence": 0.97, "_matched_ids": [],
        "date_of_purchase": "2026-05-15", "amount": 320.00, "currency": "SGD",
        "description": "Office supplies - StatioNow",
        "validation_status": "auto_validated",
    },
    {
        "invoice_id": "INV-0005", "_db_id": None, "_confidence": None, "_matched_ids": [],
        "date_of_purchase": "2026-05-18", "amount": 9800.75, "currency": "GBP",
        "description": "Legal services - Chambers & Partners",
        "validation_status": "pending",
    },
]

# Pre-seeded mock API templates used when auto-generating an endpoint for a new bank
_MOCK_API_TEMPLATE = "https://api.{slug}.morpheus.io/v1/transactions"


def _make_slug(name: str) -> str:
    """Turn a bank display name into a URL-safe slug."""
    slug = name.lower().strip()
    slug = re.sub(r"[^a-z0-9]+", "-", slug)
    return slug.strip("-")


# =============================================================================
#  HELPER WIDGETS
# =============================================================================

class Separator(ctk.CTkFrame):
    def __init__(self, master, **kwargs):
        super().__init__(master, height=1, fg_color=C["border"], **kwargs)


class StatusBadge(ctk.CTkLabel):
    def __init__(self, master, status_key, **kwargs):
        meta = STATUS_META.get(status_key, {"label": status_key, "fg": C["text_dim"]})
        super().__init__(master, text=meta["label"], text_color=meta["fg"],
                         font=ctk.CTkFont(size=12), **kwargs)

    def update_status(self, status_key):
        meta = STATUS_META.get(status_key, {"label": status_key, "fg": C["text_dim"]})
        self.configure(text=meta["label"], text_color=meta["fg"])


class StatCard(ctk.CTkFrame):
    def __init__(self, master, title, value, color=None, **kwargs):
        if color is None:
            color = C["text"]
        super().__init__(master, fg_color=C["card"], corner_radius=10,
                         border_width=1, border_color=C["border"], **kwargs)
        self.value_label = ctk.CTkLabel(self, text=value,
                                        font=ctk.CTkFont(size=26, weight="bold"),
                                        text_color=color)
        self.value_label.pack(padx=20, pady=(14, 0))
        ctk.CTkLabel(self, text=title, font=ctk.CTkFont(size=11),
                     text_color=C["text_dim"]).pack(padx=20, pady=(2, 14))

    def set_value(self, value):
        self.value_label.configure(text=value)


# =============================================================================
#  BANK MANAGER DIALOG
# =============================================================================

class BankManagerDialog(ctk.CTkToplevel):
    """Modal dialog for registering / removing banks."""

    _BANK_COL_WIDTHS = (220, 370, 70)   # name | api | action

    def __init__(self, master, db, db_rollback_fn, **kwargs):
        super().__init__(master, **kwargs)
        self.title("Registered Banks")
        self.geometry("700x540")
        self.resizable(False, False)
        self.configure(fg_color=C["bg"])

        self._db           = db
        self._db_rollback  = db_rollback_fn
        self._bank_rows    = []   # list of frame widgets currently rendered

        self._build_ui()
        self._load_banks()

        # Block interaction with the main window until this is closed
        self.grab_set()
        self.focus_set()
        self.lift()

    # ── Layout ────────────────────────────────────────────────────────────────

    def _build_ui(self):
        self.rowconfigure(0, weight=0)  # header
        self.rowconfigure(1, weight=0)  # column labels
        self.rowconfigure(2, weight=1)  # scrollable list
        self.rowconfigure(3, weight=0)  # separator
        self.rowconfigure(4, weight=0)  # add-bank form
        self.columnconfigure(0, weight=1)

        # ── Header ────────────────────────────────────────────────────────────
        hf = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0, height=58)
        hf.grid(row=0, column=0, sticky="ew")
        hf.grid_propagate(False)
        ctk.CTkLabel(hf, text="  Registered Banks",
                     font=ctk.CTkFont(size=16, weight="bold"),
                     text_color=C["text"]).pack(side="left", padx=20, pady=14)
        ctk.CTkLabel(hf,
                     text="Banks listed here are eligible for AI transaction matching.",
                     font=ctk.CTkFont(size=11), text_color=C["text_dim"]
                     ).pack(side="left", padx=0, pady=14)

        # ── Column header bar ─────────────────────────────────────────────────
        ch = ctk.CTkFrame(self, fg_color=C["card"], corner_radius=0, height=32)
        ch.grid(row=1, column=0, sticky="ew")
        ch.grid_propagate(False)
        for col, (label, width) in enumerate(zip(
                ["Bank Name", "API Endpoint", ""],
                self._BANK_COL_WIDTHS)):
            ch.columnconfigure(col, minsize=width, weight=(1 if col == 1 else 0))
            ctk.CTkLabel(ch, text=label,
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color=C["text_dim"], anchor="w"
                         ).grid(row=0, column=col, padx=(14 if col == 0 else 8),
                                pady=4, sticky="ew")

        # ── Scrollable bank list ───────────────────────────────────────────────
        self._scroll = ctk.CTkScrollableFrame(
            self, fg_color=C["bg"], corner_radius=0,
            scrollbar_button_color=C["border"],
            scrollbar_button_hover_color=C["accent"])
        self._scroll.grid(row=2, column=0, sticky="nsew")
        self._scroll.columnconfigure(0, weight=1)

        # ── Separator ─────────────────────────────────────────────────────────
        Separator(self).grid(row=3, column=0, sticky="ew", pady=(4, 0))

        # ── Add-bank form ─────────────────────────────────────────────────────
        form = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0)
        form.grid(row=4, column=0, sticky="ew", padx=0, pady=0)
        form.columnconfigure(1, weight=1)

        ctk.CTkLabel(form, text="Add Bank", font=ctk.CTkFont(size=13, weight="bold"),
                     text_color=C["text"]).grid(row=0, column=0, columnspan=4,
                                                padx=16, pady=(14, 6), sticky="w")

        # Bank name
        ctk.CTkLabel(form, text="Name", font=ctk.CTkFont(size=12),
                     text_color=C["text_dim"]).grid(row=1, column=0, padx=(16, 6),
                                                    pady=(0, 14), sticky="w")
        self._name_entry = ctk.CTkEntry(form, placeholder_text="e.g. Maybank",
                                        width=180, height=32,
                                        font=ctk.CTkFont(size=12),
                                        fg_color=C["card"], border_color=C["border"])
        self._name_entry.grid(row=1, column=1, padx=(0, 10), pady=(0, 14), sticky="w")
        self._name_entry.bind("<KeyRelease>", self._on_name_change)

        # API endpoint
        ctk.CTkLabel(form, text="API", font=ctk.CTkFont(size=12),
                     text_color=C["text_dim"]).grid(row=1, column=2, padx=(0, 6),
                                                    pady=(0, 14), sticky="w")
        api_frame = ctk.CTkFrame(form, fg_color="transparent")
        api_frame.grid(row=1, column=3, padx=(0, 16), pady=(0, 14), sticky="ew")
        form.columnconfigure(3, weight=1)

        self._api_entry = ctk.CTkEntry(api_frame,
                                       placeholder_text="https://api.yourbank.com/v1/...",
                                       height=32, font=ctk.CTkFont(size=12),
                                       fg_color=C["card"], border_color=C["border"])
        self._api_entry.pack(side="left", fill="x", expand=True, padx=(0, 8))

        ctk.CTkButton(api_frame, text="Generate",
                      width=80, height=32, font=ctk.CTkFont(size=11),
                      fg_color=C["card"], hover_color=C["border"],
                      border_width=1, border_color=C["border"],
                      command=self._generate_api).pack(side="left")

        ctk.CTkButton(api_frame, text="+ Add",
                      width=72, height=32, font=ctk.CTkFont(size=12),
                      fg_color=C["accent"], hover_color=C["accent_dim"],
                      command=self._add_bank).pack(side="left", padx=(8, 0))

    # ── Data ──────────────────────────────────────────────────────────────────

    def _load_banks(self):
        """Fetch banks from DB and re-render the list."""
        for w in self._scroll.winfo_children():
            w.destroy()
        self._bank_rows = []

        if not self._db:
            ctk.CTkLabel(self._scroll,
                         text="No database connection.",
                         font=ctk.CTkFont(size=13), text_color=C["text_muted"]
                         ).pack(pady=30)
            return

        try:
            banks = self._db.retrieve_data("registered_banks")
        except Exception:
            self._db_rollback()
            banks = []

        if not banks:
            ctk.CTkLabel(self._scroll,
                         text="No banks registered yet. Add one below.",
                         font=ctk.CTkFont(size=13), text_color=C["text_muted"]
                         ).pack(pady=30)
            return

        for idx, bank in enumerate(banks):
            bg = C["row_even"] if idx % 2 == 0 else C["row_odd"]
            row_frame = ctk.CTkFrame(self._scroll, fg_color=bg, corner_radius=0)
            row_frame.pack(fill="x", pady=0)

            for col, width in enumerate(self._BANK_COL_WIDTHS):
                row_frame.columnconfigure(col, minsize=width,
                                          weight=(1 if col == 1 else 0))

            name = bank.get("bank_name", "")
            api  = bank.get("mock_api", "") or ""

            ctk.CTkLabel(row_frame, text=name,
                         font=ctk.CTkFont(size=12, weight="bold"),
                         text_color=C["accent"], anchor="w"
                         ).grid(row=0, column=0, padx=(14, 8), pady=10, sticky="ew")

            # Truncate long URLs for display
            display_api = api if len(api) <= 52 else api[:49] + "..."
            ctk.CTkLabel(row_frame, text=display_api,
                         font=ctk.CTkFont(size=11, family="Courier"),
                         text_color=C["text_dim"], anchor="w"
                         ).grid(row=0, column=1, padx=(0, 8), pady=10, sticky="ew")

            ctk.CTkButton(row_frame, text="Remove",
                          width=64, height=26, font=ctk.CTkFont(size=11),
                          fg_color="#3a1010", hover_color="#5a1a1a",
                          text_color="#e74c3c",
                          border_width=1, border_color="#e74c3c",
                          command=lambda n=name: self._delete_bank(n)
                          ).grid(row=0, column=2, padx=(0, 10), pady=8)

            self._bank_rows.append(row_frame)

    def _add_bank(self):
        name = self._name_entry.get().strip()
        api  = self._api_entry.get().strip()

        if not name:
            messagebox.showwarning("Missing Name",
                                   "Please enter a bank name.", parent=self)
            return
        if not api:
            # Auto-generate if user didn't fill it in
            api = _MOCK_API_TEMPLATE.format(slug=_make_slug(name))

        if not self._db:
            messagebox.showerror("No Database",
                                 "Not connected to the database.", parent=self)
            return

        try:
            self._db.insert_data(
                "registered_banks",
                ["bank_name", "mock_api"],
                [name, api]
            )
        except Exception as e:
            self._db_rollback()
            err = str(e)
            if "unique" in err.lower() or "duplicate" in err.lower():
                messagebox.showerror("Duplicate",
                                     '"{}" is already registered.'.format(name),
                                     parent=self)
            else:
                messagebox.showerror("DB Error", str(e), parent=self)
            return

        self._name_entry.delete(0, "end")
        self._api_entry.delete(0, "end")
        self._load_banks()

    def _delete_bank(self, bank_name):
        if not messagebox.askyesno(
                "Confirm Removal",
                'Remove "{}" from registered banks?\n\n'
                'Future AI validations will no longer search this bank.'.format(bank_name),
                parent=self):
            return

        try:
            # Data_Retrieval has no delete_data method, so we use the raw cursor directly.
            self._db.cursor.execute(
                'DELETE FROM registered_banks WHERE bank_name = %s',
                (bank_name,)
            )
            self._db.conn.commit()
        except Exception as e:
            self._db_rollback()
            messagebox.showerror("DB Error", str(e), parent=self)
            return

        self._load_banks()

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _on_name_change(self, _event=None):
        """Keep the Generate button preview in sync as the user types."""
        pass  # generate is on-demand via button click

    def _generate_api(self):
        name = self._name_entry.get().strip()
        slug = _make_slug(name) if name else "yourbank"
        api  = _MOCK_API_TEMPLATE.format(slug=slug)
        self._api_entry.delete(0, "end")
        self._api_entry.insert(0, api)


# =============================================================================
#  ACTIVITY LOG DIALOG
# =============================================================================

_LEVEL_COLORS = {
    "info":    "#4f8ef7",
    "success": "#2ecc71",
    "warning": "#e67e22",
    "error":   "#e74c3c",
}

_LEVEL_LABELS = {
    "info":    "ℹ",
    "success": "✓",
    "warning": "⚠",
    "error":   "✗",
}


class ActivityLogDialog(ctk.CTkToplevel):
    """Non-modal window showing the in-session + persisted activity log."""

    def __init__(self, master, in_memory_log, db, db_rollback_fn, **kwargs):
        super().__init__(master, **kwargs)
        self.title("Activity Log")
        self.geometry("860x560")
        self.minsize(700, 400)
        self.configure(fg_color=C["bg"])

        self._db          = db
        self._db_rollback = db_rollback_fn
        self._entries     = []   # list of rendered entry frames

        self._build_ui()
        self._load_history(in_memory_log)

        # Non-modal — don't grab_set so main window stays usable
        self.lift()
        self.focus_set()

    # ── Layout ────────────────────────────────────────────────────────────────

    def _build_ui(self):
        self.rowconfigure(0, weight=0)   # header bar
        self.rowconfigure(1, weight=0)   # column labels
        self.rowconfigure(2, weight=1)   # scrollable log
        self.rowconfigure(3, weight=0)   # footer
        self.columnconfigure(0, weight=1)

        # ── Header ────────────────────────────────────────────────────────────
        hf = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0, height=52)
        hf.grid(row=0, column=0, sticky="ew")
        hf.grid_propagate(False)
        hf.columnconfigure(1, weight=1)

        ctk.CTkLabel(hf, text="  📋  Activity Log",
                     font=ctk.CTkFont(size=16, weight="bold"),
                     text_color=C["text"]).grid(row=0, column=0, padx=16, pady=12, sticky="w")

        ctk.CTkLabel(hf,
                     text="Live audit trail of all validation events.",
                     font=ctk.CTkFont(size=11), text_color=C["text_dim"]
                     ).grid(row=0, column=1, padx=4, pady=12, sticky="w")

        ctk.CTkButton(hf, text="Clear Log",
                      width=90, height=30, font=ctk.CTkFont(size=11),
                      fg_color="#3a1010", hover_color="#5a1a1a",
                      text_color="#e74c3c",
                      border_width=1, border_color="#e74c3c",
                      command=self._clear_log
                      ).grid(row=0, column=2, padx=16, pady=12, sticky="e")

        # ── Column labels ─────────────────────────────────────────────────────
        ch = ctk.CTkFrame(self, fg_color=C["card"], corner_radius=0, height=28)
        ch.grid(row=1, column=0, sticky="ew")
        ch.grid_propagate(False)

        col_defs = [
            ("",        32,  0),    # level dot
            ("Time",    90,  0),
            ("Invoice", 90,  0),
            ("Event",   200, 1),
            ("Details", 0,   1),
        ]
        for col_idx, (label, minw, wt) in enumerate(col_defs):
            ch.columnconfigure(col_idx, minsize=minw, weight=wt)
            ctk.CTkLabel(ch, text=label,
                         font=ctk.CTkFont(size=10, weight="bold"),
                         text_color=C["text_muted"], anchor="w"
                         ).grid(row=0, column=col_idx,
                                padx=(14 if col_idx == 0 else 6),
                                pady=3, sticky="ew")

        # ── Scrollable body ───────────────────────────────────────────────────
        self._scroll = ctk.CTkScrollableFrame(
            self, fg_color=C["bg"], corner_radius=0,
            scrollbar_button_color=C["border"],
            scrollbar_button_hover_color=C["accent"])
        self._scroll.grid(row=2, column=0, sticky="nsew")
        self._scroll.columnconfigure(0, weight=1)

        # Placeholder shown when log is empty
        self._empty_label = ctk.CTkLabel(
            self._scroll,
            text="No activity yet — run AI on an invoice to get started.",
            font=ctk.CTkFont(size=13), text_color=C["text_muted"])

        # ── Footer / status bar ───────────────────────────────────────────────
        fb = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0, height=26)
        fb.grid(row=3, column=0, sticky="ew")
        fb.grid_propagate(False)
        fb.columnconfigure(1, weight=1)

        self._count_label = ctk.CTkLabel(
            fb, text="0 entries",
            font=ctk.CTkFont(size=10), text_color=C["text_muted"])
        self._count_label.grid(row=0, column=0, padx=14, pady=3, sticky="w")

        self._db_badge = ctk.CTkLabel(
            fb,
            text=("● Persisted to DB" if self._db else "● In-memory only"),
            font=ctk.CTkFont(size=10),
            text_color="#2ecc71" if self._db else "#e67e22")
        self._db_badge.grid(row=0, column=2, padx=14, pady=3, sticky="e")

    # ── Data loading ──────────────────────────────────────────────────────────

    def _load_history(self, in_memory_log):
        """Populate the log: prefer DB rows; fall back to in-memory list."""
        entries = []

        if self._db:
            try:
                rows = self._db.retrieve_data("activity_log")
                # retrieve_data returns a list; convert to our entry format
                for r in (rows or []):
                    entries.append({
                        "time":        r.get("logged_at") or datetime.now(),
                        "invoice_ref": r.get("invoice_ref") or "",
                        "action":      r.get("action") or "",
                        "details":     r.get("details") or "",
                        "level":       r.get("level") or "info",
                    })
            except Exception:
                self._db_rollback()
                entries = list(in_memory_log)
        else:
            entries = list(in_memory_log)

        if not entries:
            self._show_empty()
            return

        for entry in entries:
            self._render_entry(entry)
        self._scroll_to_bottom()

    def _show_empty(self):
        self._empty_label.pack(pady=40)
        self._count_label.configure(text="0 entries")

    def _hide_empty(self):
        self._empty_label.pack_forget()

    # ── Public API ────────────────────────────────────────────────────────────

    def append_entry(self, entry):
        """Called live by _log_activity when the dialog is already open."""
        self._hide_empty()
        self._render_entry(entry)
        self._scroll_to_bottom()

    # ── Rendering ─────────────────────────────────────────────────────────────

    def _render_entry(self, entry):
        idx = len(self._entries)
        bg  = C["row_even"] if idx % 2 == 0 else C["row_odd"]

        level   = entry.get("level", "info")
        color   = _LEVEL_COLORS.get(level, C["text_dim"])
        icon    = _LEVEL_LABELS.get(level, "•")
        ts      = entry.get("time")
        ts_str  = ts.strftime("%H:%M:%S") if hasattr(ts, "strftime") else str(ts)[:8]
        ref     = entry.get("invoice_ref") or "—"
        action  = entry.get("action") or ""
        details = entry.get("details") or ""

        row = ctk.CTkFrame(self._scroll, fg_color=bg, corner_radius=0)
        row.pack(fill="x", pady=0)
        row.columnconfigure(3, weight=1)
        row.columnconfigure(4, weight=2)

        # Level icon
        ctk.CTkLabel(row, text=icon, font=ctk.CTkFont(size=13, weight="bold"),
                     text_color=color, width=28, anchor="center"
                     ).grid(row=0, column=0, padx=(10, 2), pady=6)

        # Time
        ctk.CTkLabel(row, text=ts_str, font=ctk.CTkFont(size=11, family="Courier"),
                     text_color=C["text_muted"], anchor="w", width=72
                     ).grid(row=0, column=1, padx=(2, 6), pady=6, sticky="w")

        # Invoice ref
        ctk.CTkLabel(row, text=ref, font=ctk.CTkFont(size=11, weight="bold"),
                     text_color=color, anchor="w", width=82
                     ).grid(row=0, column=2, padx=(0, 6), pady=6, sticky="w")

        # Action
        ctk.CTkLabel(row, text=action, font=ctk.CTkFont(size=12),
                     text_color=C["text"], anchor="w"
                     ).grid(row=0, column=3, padx=(0, 8), pady=6, sticky="ew")

        # Details (dimmed, may be long — truncate gracefully)
        if details:
            disp = details if len(details) <= 80 else details[:77] + "..."
            ctk.CTkLabel(row, text=disp, font=ctk.CTkFont(size=11),
                         text_color=C["text_dim"], anchor="w",
                         wraplength=300
                         ).grid(row=0, column=4, padx=(0, 12), pady=6, sticky="ew")

        self._entries.append(row)
        self._count_label.configure(text="{} entr{}".format(
            len(self._entries), "y" if len(self._entries) == 1 else "ies"))

    def _scroll_to_bottom(self):
        """Push the scrollable canvas to the bottom so newest entry is visible."""
        try:
            self._scroll._parent_canvas.yview_moveto(1.0)
        except Exception:
            pass

    # ── Clear ─────────────────────────────────────────────────────────────────

    def _clear_log(self):
        if not messagebox.askyesno(
                "Clear Log",
                "Delete all activity log entries?\n\n"
                "This will also remove them from the database.",
                parent=self):
            return

        # Remove from DB
        if self._db:
            try:
                self._db.cursor.execute("DELETE FROM activity_log")
                self._db.conn.commit()
            except Exception:
                self._db_rollback()

        # Remove in-memory entries from the parent app list
        if hasattr(self.master, "_activity_log"):
            self.master._activity_log.clear()

        # Destroy rendered rows
        for w in self._entries:
            try:
                w.destroy()
            except Exception:
                pass
        self._entries.clear()
        self._show_empty()


# =============================================================================
#  INVOICE ROW
# =============================================================================

class InvoiceRow(ctk.CTkFrame):
    def __init__(self, master, invoice, row_index,
                 on_run_ai, on_approve, on_select_change, **kwargs):
        bg = C["row_even"] if row_index % 2 == 0 else C["row_odd"]
        super().__init__(master, fg_color=bg, corner_radius=0, **kwargs)

        self.invoice    = invoice
        self.row_index  = row_index
        self.on_run_ai  = on_run_ai
        self.on_approve = on_approve
        self._bg        = bg
        self._selected  = tk.BooleanVar(value=False)

        self.bind("<Enter>", self._on_enter)
        self.bind("<Leave>", self._on_leave)
        self._build(on_select_change)

    def _build(self, on_select_change):
        for i, (key, _label, width, _anchor) in enumerate(COLUMNS):
            self.columnconfigure(i, minsize=width,
                                 weight=(1 if key == "description" else 0))
        inv = self.invoice
        col = 0

        # Checkbox
        ctk.CTkCheckBox(self, variable=self._selected, text="",
                        width=24, height=24, command=on_select_change,
                        fg_color=C["accent"], hover_color=C["accent_dim"],
                        border_color=C["border"]
                        ).grid(row=0, column=col, padx=(8, 4), pady=8)
        col += 1

        # Invoice ID
        ctk.CTkLabel(self, text=inv.get("invoice_id", "-"),
                     font=ctk.CTkFont(size=12, weight="bold"),
                     text_color=C["accent"], anchor="w"
                     ).grid(row=0, column=col, padx=(4, 8), pady=8, sticky="ew")
        col += 1

        # Date
        raw_date = inv.get("date_of_purchase", "")
        try:
            date_str = datetime.strptime(str(raw_date), "%Y-%m-%d").strftime("%d %b %Y")
        except (ValueError, TypeError):
            date_str = str(raw_date) if raw_date else "-"
        ctk.CTkLabel(self, text=date_str, font=ctk.CTkFont(size=12),
                     text_color=C["text_dim"], anchor="center"
                     ).grid(row=0, column=col, padx=4, pady=8, sticky="ew")
        col += 1

        # Amount
        try:
            amt = float(inv.get("amount", 0))
            amount_str = "{:,.2f}".format(amt) if amt != 0 else "-"
        except (ValueError, TypeError):
            amount_str = "-"
        ctk.CTkLabel(self, text=amount_str,
                     font=ctk.CTkFont(size=12, family="Courier"),
                     text_color=C["text"], anchor="e"
                     ).grid(row=0, column=col, padx=(4, 2), pady=8, sticky="ew")
        col += 1

        # Currency
        ctk.CTkLabel(self, text=inv.get("currency", ""),
                     font=ctk.CTkFont(size=11), text_color=C["text_muted"], anchor="center"
                     ).grid(row=0, column=col, padx=(2, 4), pady=8, sticky="ew")
        col += 1

        # Description
        self.desc_label = ctk.CTkLabel(self, text=inv.get("description", ""),
                                       font=ctk.CTkFont(size=12),
                                       text_color=C["text"], anchor="w")
        self.desc_label.grid(row=0, column=col, padx=(4, 8), pady=8, sticky="ew")
        col += 1

        # Confidence
        conf = inv.get("_confidence")
        if conf is not None:
            color = "#2ecc71" if conf >= CONFIDENCE_THRESHOLD else "#e67e22"
            conf_text = "{:.1f}%".format(conf * 100)
        else:
            color = C["text_muted"]
            conf_text = "-"
        self.confidence_label = ctk.CTkLabel(self, text=conf_text,
                                              font=ctk.CTkFont(size=12, weight="bold"),
                                              text_color=color, anchor="center")
        self.confidence_label.grid(row=0, column=col, padx=4, pady=8, sticky="ew")
        col += 1

        # Status badge
        self.status_badge = StatusBadge(self,
                                         status_key=inv.get("validation_status", "pending"),
                                         anchor="center")
        self.status_badge.grid(row=0, column=col, padx=4, pady=8, sticky="ew")
        col += 1

        # Action button area
        self.action_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.action_frame.grid(row=0, column=col, padx=8, pady=6)
        self._refresh_action_button()

    def _refresh_action_button(self):
        for w in self.action_frame.winfo_children():
            w.destroy()
        status = self.invoice.get("validation_status", "pending")

        if status == "pending":
            ctk.CTkButton(self.action_frame, text="Run AI",
                          width=110, height=30, font=ctk.CTkFont(size=12),
                          fg_color=C["accent"], hover_color=C["accent_dim"],
                          command=lambda: self.on_run_ai(self)).pack()

        elif status == "needs_manual":
            ctk.CTkButton(self.action_frame, text="Approve",
                          width=110, height=30, font=ctk.CTkFont(size=12),
                          fg_color="#1a4a1a", hover_color="#265526",
                          text_color="#2ecc71",
                          border_width=1, border_color="#2ecc71",
                          command=lambda: self.on_approve(self)).pack()

        elif status == "processing":
            ctk.CTkLabel(self.action_frame, text="Running...",
                         font=ctk.CTkFont(size=12), text_color=C["accent"]).pack()

        elif status == "error":
            # Show a Details popup (with the exact error reason) + a Retry button
            msg = self.invoice.get("_message", "Unknown error — check the activity log.")
            ctk.CTkButton(self.action_frame, text="Details",
                          width=62, height=30, font=ctk.CTkFont(size=11),
                          fg_color="#2a1010", hover_color="#3e1515",
                          text_color="#e74c3c",
                          border_width=1, border_color="#e74c3c",
                          command=lambda m=msg: tk.messagebox.showinfo(
                              "Validation Error", m)
                          ).pack(side="left", padx=(0, 4))
            ctk.CTkButton(self.action_frame, text="Retry",
                          width=42, height=30, font=ctk.CTkFont(size=11),
                          fg_color=C["card"], hover_color=C["border"],
                          border_width=1, border_color=C["border"],
                          command=lambda: self.on_run_ai(self)).pack(side="left")

    def set_processing(self):
        self.invoice["validation_status"] = "processing"
        self.status_badge.update_status("processing")
        self.confidence_label.configure(text="...", text_color=C["accent"])
        self._refresh_action_button()

    def set_result(self, status, confidence, matched_ids=None, message=""):
        self.invoice["validation_status"] = status
        self.invoice["_message"] = message   # store for Details button
        self.invoice["_confidence"] = confidence
        self.invoice["_matched_ids"] = matched_ids or []
        self.status_badge.update_status(status)
        if confidence is not None:
            pct   = confidence * 100
            color = "#2ecc71" if confidence >= CONFIDENCE_THRESHOLD else "#e67e22"
            self.confidence_label.configure(text="{:.1f}%".format(pct), text_color=color)
        else:
            self.confidence_label.configure(text="-", text_color=C["text_muted"])
        self._refresh_action_button()

    def set_approved(self):
        self.invoice["validation_status"] = "manually_validated"
        self.status_badge.update_status("manually_validated")
        self._refresh_action_button()

    def _on_enter(self, _event=None):
        self.configure(fg_color=C["row_sel"])

    def _on_leave(self, _event=None):
        self.configure(fg_color=self._bg)

    @property
    def is_selected(self):
        return self._selected.get()

    def set_selected(self, value):
        self._selected.set(value)


# =============================================================================
#  MANUAL VALIDATION REVIEW
# =============================================================================

class ManualValidationDialog(ctk.CTkToplevel):
    """Read-only cross-check window shown before a manual validation is approved."""

    def __init__(self, master, invoice, review_data, on_approve, **kwargs):
        super().__init__(master, **kwargs)
        self.title("Manual Validation Review")
        self.geometry("920x680")
        self.minsize(760, 560)
        self.configure(fg_color=C["bg"])

        self._invoice = invoice
        self._review_data = review_data or {}
        self._on_approve = on_approve
        self._preview_images = []

        self._build_ui()
        self.grab_set()
        self.focus_set()
        self.lift()

    def _build_ui(self):
        self.rowconfigure(1, weight=1)
        self.columnconfigure(0, weight=1)

        header = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0)
        header.grid(row=0, column=0, sticky="ew")
        header.columnconfigure(0, weight=1)

        title = "Manual validation: {}".format(self._invoice.get("invoice_id", "-"))
        ctk.CTkLabel(header, text=title,
                     font=ctk.CTkFont(size=18, weight="bold"),
                     text_color=C["text"], anchor="w"
                     ).grid(row=0, column=0, padx=20, pady=(16, 2), sticky="ew")

        conf = self._invoice.get("_confidence")
        conf_text = "Confidence: {:.1f}%".format(conf * 100) if conf is not None else "Confidence: -"
        ctk.CTkLabel(header,
                     text="Cross-check the database values before approving. {}".format(conf_text),
                     font=ctk.CTkFont(size=12),
                     text_color=C["text_dim"], anchor="w"
                     ).grid(row=1, column=0, padx=20, pady=(0, 16), sticky="ew")

        body = ctk.CTkScrollableFrame(self, fg_color=C["bg"], corner_radius=0)
        body.grid(row=1, column=0, sticky="nsew", padx=18, pady=18)
        body.columnconfigure(0, weight=1)
        body.columnconfigure(1, weight=1)

        self._section(body, "Invoice Summary", self._invoice_summary_items(), 0, 0)
        if self._should_show_source_file():
            self._file_preview_section(body, 0, 1)
        else:
            self._section(body, "Extracted OCR Data", self._dict_items(self._review_data.get("ocr_result")), 0, 1)
        self._section(body, "Invoice DB Row", self._dict_items(self._review_data.get("invoice_row")), 1, 0)
        self._transactions_section(body, 1, 1)

        footer = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0)
        footer.grid(row=2, column=0, sticky="ew")
        footer.columnconfigure(0, weight=1)

        ctk.CTkButton(footer, text="Close",
                      width=120, height=34,
                      fg_color=C["card"], hover_color=C["border"],
                      border_width=1, border_color=C["border"],
                      command=self.destroy
                      ).grid(row=0, column=1, padx=(0, 10), pady=14)
        ctk.CTkButton(footer, text="Approve Validation",
                      width=170, height=34,
                      fg_color="#1a4a1a", hover_color="#265526",
                      text_color="#2ecc71",
                      border_width=1, border_color="#2ecc71",
                      command=self._approve
                      ).grid(row=0, column=2, padx=(0, 18), pady=14)

    def _section(self, master, title, items, row, col):
        frame = ctk.CTkFrame(master, fg_color=C["card"], corner_radius=8,
                             border_width=1, border_color=C["border"])
        frame.grid(row=row, column=col, sticky="nsew", padx=8, pady=8)
        frame.columnconfigure(1, weight=1)

        ctk.CTkLabel(frame, text=title,
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color=C["text"], anchor="w"
                     ).grid(row=0, column=0, columnspan=2, padx=14, pady=(12, 8), sticky="ew")

        if not items:
            ctk.CTkLabel(frame, text="No data found.",
                         font=ctk.CTkFont(size=12), text_color=C["text_muted"],
                         anchor="w").grid(row=1, column=0, columnspan=2, padx=14, pady=(0, 14), sticky="ew")
            return

        for idx, (key, value) in enumerate(items, start=1):
            ctk.CTkLabel(frame, text=str(key),
                         font=ctk.CTkFont(size=11, weight="bold"),
                         text_color=C["text_dim"], anchor="w"
                         ).grid(row=idx, column=0, padx=(14, 10), pady=3, sticky="nw")
            ctk.CTkLabel(frame, text=self._display_value(value),
                         font=ctk.CTkFont(size=12),
                         text_color=C["text"], anchor="w",
                         wraplength=260, justify="left"
                         ).grid(row=idx, column=1, padx=(0, 14), pady=3, sticky="ew")

    def _transactions_section(self, master, row, col):
        txns = self._review_data.get("transactions") or []
        items = []
        for idx, txn in enumerate(txns, start=1):
            if isinstance(txn, dict):
                txn_id = txn.get("transaction_id") or txn.get("transaction_ID") or "Transaction {}".format(idx)
                detail_bits = []
                for key in ("bank_name", "transaction_datetime", "amount", "currency", "description"):
                    if key in txn and txn.get(key) not in (None, ""):
                        detail_bits.append("{}: {}".format(key, self._display_value(txn.get(key))))
                items.append((txn_id, "\n".join(detail_bits) if detail_bits else txn))
            else:
                items.append(("Transaction {}".format(idx), txn))
        self._section(master, "Matched Transaction Data", items, row, col)

    def _file_preview_section(self, master, row, col):
        info = self._source_file_info()
        path = info.get("file_path", "")
        file_type = (info.get("file_type") or "").lower()
        title = "Original Invoice File"

        frame = ctk.CTkFrame(master, fg_color=C["card"], corner_radius=8,
                             border_width=1, border_color=C["border"])
        frame.grid(row=row, column=col, sticky="nsew", padx=8, pady=8)
        frame.columnconfigure(0, weight=1)

        ctk.CTkLabel(frame, text=title,
                     font=ctk.CTkFont(size=14, weight="bold"),
                     text_color=C["text"], anchor="w"
                     ).grid(row=0, column=0, padx=14, pady=(12, 4), sticky="ew")

        ctk.CTkLabel(frame,
                     text="{} file from database path:\n{}".format(file_type.upper(), path or "-"),
                     font=ctk.CTkFont(size=11),
                     text_color=C["text_dim"], anchor="w",
                     wraplength=340, justify="left"
                     ).grid(row=1, column=0, padx=14, pady=(0, 8), sticky="ew")

        ctk.CTkButton(frame, text="Open Source File",
                      width=150, height=30,
                      fg_color=C["accent"], hover_color=C["accent_dim"],
                      command=lambda p=path: self._open_source_file(p)
                      ).grid(row=2, column=0, padx=14, pady=(0, 10), sticky="w")

        if not path or not os.path.exists(path):
            self._preview_message(frame, "File not found at the stored database path.", 3)
            return

        if file_type == "pdf":
            photo, error = self._render_pdf_preview(path)
            if photo:
                self._preview_images.append(photo)
                tk.Label(frame, image=photo, bg=C["card"], borderwidth=0
                         ).grid(row=3, column=0, padx=14, pady=(0, 14), sticky="n")
            else:
                self._preview_message(frame, error or "PDF preview unavailable.", 3)
        elif file_type == "docx":
            text, error = self._extract_docx_preview(path)
            if text:
                box = ctk.CTkTextbox(frame, height=260, fg_color=C["row_even"],
                                     text_color=C["text"], border_width=1,
                                     border_color=C["border"], wrap="word")
                box.grid(row=3, column=0, padx=14, pady=(0, 14), sticky="nsew")
                box.insert("1.0", text)
                box.configure(state="disabled")
            else:
                self._preview_message(frame, error or "DOCX preview unavailable.", 3)
        else:
            self._preview_message(frame, "Preview is only enabled for PDF and DOCX invoices.", 3)

    def _preview_message(self, frame, message, row):
        ctk.CTkLabel(frame, text=message,
                     font=ctk.CTkFont(size=12),
                     text_color=C["text_muted"], anchor="w",
                     wraplength=340, justify="left"
                     ).grid(row=row, column=0, padx=14, pady=(0, 14), sticky="ew")

    def _source_file_info(self):
        row = self._review_data.get("invoice_row") or {}
        return {
            "file_path": row.get("file_path") or row.get("File_Path") or "",
            "file_type": row.get("file_type") or row.get("File_Type") or "",
        }

    def _should_show_source_file(self):
        info = self._source_file_info()
        return (info.get("file_type") or "").lower() in ("pdf", "docx")

    def _render_pdf_preview(self, path):
        try:
            import fitz
            from PIL import Image, ImageTk

            doc = fitz.open(path)
            try:
                page = doc.load_page(0)
                pix = page.get_pixmap(matrix=fitz.Matrix(1.25, 1.25), alpha=False)
                image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
                image.thumbnail((360, 460))
                return ImageTk.PhotoImage(image), None
            finally:
                doc.close()
        except Exception as e:
            return None, "Could not render PDF preview: {}".format(e)

    def _extract_docx_preview(self, path):
        try:
            from docx import Document

            doc = Document(path)
            parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for table_row in table.rows:
                    cells = [cell.text.strip() for cell in table_row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            text = "\n".join(parts).strip()
            if len(text) > 4000:
                text = text[:4000] + "\n\n[Preview truncated]"
            return text, None
        except Exception as e:
            return "", "Could not read DOCX preview: {}".format(e)

    def _open_source_file(self, path):
        if not path or not os.path.exists(path):
            messagebox.showerror("File Not Found",
                                 "The source file does not exist at the stored path:\n{}".format(path),
                                 parent=self)
            return
        try:
            os.startfile(path)
        except Exception as e:
            messagebox.showerror("Open File Failed", str(e), parent=self)

    def _invoice_summary_items(self):
        fields = [
            ("Invoice ID", self._invoice.get("invoice_id")),
            ("Date", self._invoice.get("date_of_purchase")),
            ("Amount", self._invoice.get("amount")),
            ("Currency", self._invoice.get("currency")),
            ("Description", self._invoice.get("description")),
            ("Matched IDs", ", ".join(str(t) for t in self._invoice.get("_matched_ids", []) or [])),
        ]
        return [(k, v) for k, v in fields if v not in (None, "")]

    def _dict_items(self, value):
        if not value:
            return []
        if isinstance(value, dict):
            return list(value.items())
        return [("Value", value)]

    def _display_value(self, value):
        if value is None:
            return "-"
        if isinstance(value, float):
            return "{:,.4f}".format(value).rstrip("0").rstrip(".")
        return str(value)

    def _approve(self):
        self.destroy()
        self._on_approve()


# =============================================================================
#  MAIN APPLICATION
# =============================================================================

class GlobalTreasuryApp(ctk.CTk):

    def __init__(self):
        super().__init__()
        self.title("Global Treasury Agent")
        self.geometry("1280x780")
        self.minsize(1100, 600)
        self.configure(fg_color=C["bg"])

        self._invoices       = []
        self._rows           = []
        self._db             = None
        self._agent          = None
        self._active_threads = {}
        self._filter_var     = None
        self._search_entry   = None
        self._bank_dialog    = None   # keep reference to avoid GC
        self._log_dialog     = None   # activity log window
        self._manual_review_dialogs = {}
        self._activity_log   = []     # in-memory log entries

        self._build_ui()
        self._connect_backend()
        self._init_activity_log()     # auto-create DB table if needed
        self._load_invoices()

    # =========================================================================
    #  UI CONSTRUCTION
    # =========================================================================

    def _build_ui(self):
        self.rowconfigure(0, weight=0)
        self.rowconfigure(1, weight=0)
        self.rowconfigure(2, weight=0)
        self.rowconfigure(3, weight=0)
        self.rowconfigure(4, weight=1)
        self.rowconfigure(5, weight=0)
        self.columnconfigure(0, weight=1)

        self._build_header()
        self._build_toolbar()
        Separator(self).grid(row=2, column=0, sticky="ew")
        self._build_table_header()
        self._build_table_body()
        self._build_status_bar()

    def _build_header(self):
        hf = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0, height=90)
        hf.grid(row=0, column=0, sticky="ew")
        hf.grid_propagate(False)
        hf.columnconfigure(1, weight=1)

        title_frame = ctk.CTkFrame(hf, fg_color="transparent")
        title_frame.grid(row=0, column=0, padx=24, pady=16, sticky="w")
        ctk.CTkLabel(title_frame, text="  Global Treasury Agent",
                     font=ctk.CTkFont(size=20, weight="bold"),
                     text_color=C["text"]).pack(anchor="w")
        ctk.CTkLabel(title_frame, text="  Invoice Validation Dashboard",
                     font=ctk.CTkFont(size=12),
                     text_color=C["text_dim"]).pack(anchor="w")

        cards_frame = ctk.CTkFrame(hf, fg_color="transparent")
        cards_frame.grid(row=0, column=2, padx=24, pady=12, sticky="e")

        self._card_total   = StatCard(cards_frame, "Total Invoices", "-", C["text"])
        self._card_pending = StatCard(cards_frame, "Pending",        "-", "#f5a623")
        self._card_review  = StatCard(cards_frame, "Needs Review",   "-", "#e67e22")
        self._card_done    = StatCard(cards_frame, "Validated",      "-", "#2ecc71")

        for i, card in enumerate([self._card_total, self._card_pending,
                                   self._card_review, self._card_done]):
            card.grid(row=0, column=i, padx=6)

        self._db_dot = ctk.CTkLabel(hf, text="●", font=ctk.CTkFont(size=14),
                                    text_color=C["text_muted"])
        self._db_dot.grid(row=0, column=2, padx=(0, 12), pady=8, sticky="ne")

    def _build_toolbar(self):
        tf = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0, height=54)
        tf.grid(row=1, column=0, sticky="ew")
        tf.grid_propagate(False)
        tf.columnconfigure(1, weight=1)

        # Search
        search_frame = ctk.CTkFrame(tf, fg_color="transparent")
        search_frame.grid(row=0, column=0, padx=(16, 8), pady=10, sticky="w")
        self._search_entry = ctk.CTkEntry(search_frame,
                                          placeholder_text="Search invoices...",
                                          width=220, height=32,
                                          font=ctk.CTkFont(size=12),
                                          fg_color=C["card"],
                                          border_color=C["border"])
        self._search_entry.pack(side="left")
        self._search_entry.bind("<KeyRelease>", self._on_search)

        # Filter dropdown
        self._filter_var = ctk.StringVar(value="Show: All")
        ctk.CTkOptionMenu(tf, variable=self._filter_var,
                          values=["Show: All", "Show: Pending",
                                  "Show: Needs Review", "Show: Validated"],
                          width=160, height=32, font=ctk.CTkFont(size=12),
                          fg_color=C["card"], button_color=C["border"],
                          button_hover_color=C["accent"],
                          command=self._on_filter_change
                          ).grid(row=0, column=1, padx=8, pady=10, sticky="w")

        # Right-side buttons
        btn_frame = ctk.CTkFrame(tf, fg_color="transparent")
        btn_frame.grid(row=0, column=2, padx=(8, 16), pady=10, sticky="e")

        self._select_all_var = tk.BooleanVar(value=False)
        ctk.CTkCheckBox(btn_frame, text="Select All",
                        variable=self._select_all_var,
                        font=ctk.CTkFont(size=12), text_color=C["text_dim"],
                        fg_color=C["accent"], hover_color=C["accent_dim"],
                        border_color=C["border"],
                        command=self._on_select_all).pack(side="left", padx=(0, 12))

        ctk.CTkButton(btn_frame, text="Run AI on Selected",
                      width=160, height=32, font=ctk.CTkFont(size=12),
                      fg_color=C["accent"], hover_color=C["accent_dim"],
                      command=self._run_selected).pack(side="left", padx=(0, 8))

        # Upload button (green-tinted to signal "add" action)
        ctk.CTkButton(btn_frame, text="+ Upload Invoice",
                      width=140, height=32, font=ctk.CTkFont(size=12),
                      fg_color="#1a4a2e", hover_color="#235e3a",
                      text_color="#2ecc71",
                      border_width=1, border_color="#2ecc71",
                      command=self._upload_invoices).pack(side="left", padx=(0, 8))

        # ── Bank Manager button ───────────────────────────────────────────────
        ctk.CTkButton(btn_frame, text="🏦 Banks",
                      width=90, height=32, font=ctk.CTkFont(size=12),
                      fg_color=C["card"], hover_color=C["border"],
                      border_width=1, border_color=C["border"],
                      command=self._open_bank_manager).pack(side="left", padx=(0, 8))

        # ── Activity Log button ───────────────────────────────────────────────
        ctk.CTkButton(btn_frame, text="📋 Log",
                      width=76, height=32, font=ctk.CTkFont(size=12),
                      fg_color=C["card"], hover_color=C["border"],
                      border_width=1, border_color=C["border"],
                      command=self._open_activity_log).pack(side="left", padx=(0, 8))

        ctk.CTkButton(btn_frame, text="Refresh",
                      width=100, height=32, font=ctk.CTkFont(size=12),
                      fg_color=C["card"], hover_color=C["border"],
                      border_width=1, border_color=C["border"],
                      command=self._refresh).pack(side="left")

    def _build_table_header(self):
        hf = ctk.CTkFrame(self, fg_color=C["card"], corner_radius=0, height=36)
        hf.grid(row=3, column=0, sticky="ew")
        hf.grid_propagate(False)
        for i, (key, label, width, anchor) in enumerate(COLUMNS):
            hf.columnconfigure(i, minsize=width,
                               weight=(1 if key == "description" else 0))
            ctk.CTkLabel(hf, text=label, font=ctk.CTkFont(size=11, weight="bold"),
                         text_color=C["text_dim"], anchor=anchor
                         ).grid(row=0, column=i,
                                padx=(8 if i == 0 else 4), pady=4, sticky="ew")

    def _build_table_body(self):
        self._scroll = ctk.CTkScrollableFrame(
            self, fg_color=C["bg"], corner_radius=0,
            scrollbar_button_color=C["border"],
            scrollbar_button_hover_color=C["accent"])
        self._scroll.grid(row=4, column=0, sticky="nsew")
        self._scroll.columnconfigure(0, weight=1)

    def _build_status_bar(self):
        sb = ctk.CTkFrame(self, fg_color=C["panel"], corner_radius=0, height=28)
        sb.grid(row=5, column=0, sticky="ew")
        sb.grid_propagate(False)
        sb.columnconfigure(1, weight=1)
        self._db_status_label = ctk.CTkLabel(sb, text="",
                                              font=ctk.CTkFont(size=11),
                                              text_color=C["text_dim"])
        self._db_status_label.grid(row=0, column=0, padx=14, pady=4, sticky="w")
        self._last_refresh_label = ctk.CTkLabel(sb, text="",
                                                 font=ctk.CTkFont(size=11),
                                                 text_color=C["text_muted"])
        self._last_refresh_label.grid(row=0, column=2, padx=14, pady=4, sticky="e")

    # =========================================================================
    #  BACKEND CONNECTION
    # =========================================================================

    def _connect_backend(self):
        if _DB_AVAILABLE:
            try:
                self._db = Database_Connector()
                self._db_dot.configure(text_color="#2ecc71")
                self._db_status_label.configure(text="● Database connected")
            except Exception as e:
                self._db = None
                self._db_dot.configure(text_color="#e74c3c")
                self._db_status_label.configure(text="● DB error: {}".format(e))
        else:
            self._db_dot.configure(text_color="#f5a623")
            self._db_status_label.configure(
                text="● Database module unavailable - using mock data")

        if _AGENT_AVAILABLE and self._db:
            try:
                self._agent = Agent(self._db)
            except Exception:
                self._agent = None

    # =========================================================================
    #  DATA LOADING
    # =========================================================================

    def _db_rollback(self):
        """Reset a failed psycopg2 transaction so the connection can be reused."""
        try:
            self._db.conn.rollback()
        except Exception:
            pass

    def _load_invoices(self):
        if self._db:
            # Always rollback first to clear any aborted-transaction state
            self._db_rollback()
            try:
                raw = self._db.retrieve_data("invoices")

                import json as _json

                # Build OCR lookup: db_id (int) -> parsed JSONB dict
                ocr_lookup = {}
                try:
                    for ocr_row in self._db.retrieve_data("ocr_results"):
                        db_id  = ocr_row.get("invoice_id") or ocr_row.get("invoice_ID")
                        result = ocr_row.get("ocr_result") or ocr_row.get("OCR_result")
                        if db_id is not None and result:
                            if isinstance(result, str):
                                result = _json.loads(result)
                            ocr_lookup[int(db_id)] = result
                except Exception:
                    self._db_rollback()

                # confidence lookup: db_id -> confidence_score (0-1 float)
                conf_lookup = {}
                try:
                    for val_row in self._db.retrieve_data("validation_details"):
                        did = val_row.get("invoice_id")
                        if did is not None:
                            conf_lookup[int(did)] = float(val_row.get("confidence_score", 0))
                except Exception:
                    self._db_rollback()

                # matched transaction IDs: db_id -> [transaction_id, ...]
                txn_lookup = {}
                try:
                    for vt_row in self._db.retrieve_data("validation_transactions"):
                        did = vt_row.get("invoice_id")
                        tid = vt_row.get("transaction_id")
                        if did is not None and tid is not None:
                            txn_lookup.setdefault(int(did), []).append(str(tid))
                except Exception:
                    self._db_rollback()

                self._invoices = []
                for r in raw:
                    db_id = int(r.get("invoice_id") or r.get("invoice_ID", 0))
                    ocr   = ocr_lookup.get(db_id, {})

                    validated  = bool(r.get("validation_status", False))
                    confidence = conf_lookup.get(db_id)
                    matched    = txn_lookup.get(db_id, [])

                    if validated:
                        if confidence is not None and confidence >= CONFIDENCE_THRESHOLD:
                            ui_status = "auto_validated"
                        else:
                            ui_status = "manually_validated"
                    else:
                        ui_status  = "pending"
                        confidence = None

                    self._invoices.append({
                        "invoice_id":        "INV-{:04d}".format(db_id),
                        "_db_id":            db_id,
                        "_confidence":       confidence,
                        "_matched_ids":      matched,
                        "date_of_purchase":  str(ocr.get("date", "")),
                        "amount":            float(ocr.get("invoice_amount", 0) or 0),
                        "currency":          str(ocr.get("currency", "")),
                        "description":       ocr.get("vendor") or str(r.get("file_name", "")),
                        "validation_status": ui_status,
                    })

            except Exception as e:
                self._db_rollback()
                messagebox.showerror("Load Error", "Could not load invoices:\n{}".format(e))
                self._invoices = list(MOCK_INVOICES)
        else:
            self._invoices = list(MOCK_INVOICES)

        self._render_invoices()
        self._update_stats()
        self._last_refresh_label.configure(
            text="Last refreshed: {}".format(datetime.now().strftime("%H:%M:%S")))

    def _render_invoices(self):
        for w in self._scroll.winfo_children():
            w.destroy()
        self._rows = []

        filtered = self._apply_filters(self._invoices)

        if not filtered:
            ctk.CTkLabel(self._scroll, text="No invoices match the current filter.",
                         font=ctk.CTkFont(size=14),
                         text_color=C["text_muted"]).pack(pady=60)
            return

        for idx, inv in enumerate(filtered):
            row = InvoiceRow(self._scroll, inv, idx,
                             on_run_ai=self._run_ai_for_row,
                             on_approve=self._approve_row,
                             on_select_change=self._update_select_all_state)
            row.grid(row=idx * 2, column=0, sticky="ew")
            self._scroll.columnconfigure(0, weight=1)
            self._rows.append(row)
            if idx < len(filtered) - 1:
                Separator(self._scroll).grid(row=idx * 2 + 1, column=0, sticky="ew")

    def _apply_filters(self, invoices):
        fv    = self._filter_var.get() if self._filter_var else "Show: All"
        query = self._search_entry.get().lower().strip() if self._search_entry else ""

        result = []
        for inv in invoices:
            s = inv.get("validation_status", "")
            if fv == "Show: Pending"      and s != "pending":          continue
            if fv == "Show: Needs Review" and s != "needs_manual":     continue
            if fv == "Show: Validated"    and s not in ("auto_validated", "manually_validated"): continue
            if query:
                haystack = " ".join([
                    inv.get("invoice_id", ""),
                    inv.get("description", ""),
                    inv.get("currency", ""),
                    str(inv.get("amount", "")),
                ]).lower()
                if query not in haystack:
                    continue
            result.append(inv)
        return result

    # =========================================================================
    #  STATS
    # =========================================================================

    def _update_stats(self):
        total   = len(self._invoices)
        pending = sum(1 for i in self._invoices if i["validation_status"] == "pending")
        review  = sum(1 for i in self._invoices if i["validation_status"] == "needs_manual")
        done    = sum(1 for i in self._invoices
                      if i["validation_status"] in ("auto_validated", "manually_validated"))
        self._card_total.set_value(str(total))
        self._card_pending.set_value(str(pending))
        self._card_review.set_value(str(review))
        self._card_done.set_value(str(done))

    # =========================================================================
    #  UPLOAD
    # =========================================================================

    def _upload_invoices(self):
        if not self._db:
            messagebox.showerror("No Database",
                                 "Cannot upload: not connected to the database.\n"
                                 "Running in mock mode.")
            return

        paths = filedialog.askopenfilenames(
            title="Select Invoice Files",
            filetypes=[
                ("Invoice files", "*.pdf *.docx *.jpg *.jpeg *.png"),
                ("PDF",   "*.pdf"),
                ("Word",  "*.docx"),
                ("Image", "*.jpg *.jpeg *.png"),
            ]
        )
        if not paths:
            return

        uploaded, skipped, errors = [], [], []

        for path in paths:
            filename = os.path.basename(path)
            ext      = os.path.splitext(filename)[1].lower().lstrip(".")

            if ext not in UPLOAD_EXT_MAP:
                skipped.append("{} (unsupported type .{})".format(filename, ext))
                continue

            file_type, requires_ocr = UPLOAD_EXT_MAP[ext]

            try:
                self._db.insert_data(
                    "invoices",
                    ["file_name", "file_path", "file_type", "requires_ocr", "validation_status"],
                    [filename, path, file_type, requires_ocr, False]
                )
                uploaded.append(filename)
                self._log_activity(
                    filename, "Invoice Uploaded",
                    "Type: {} | OCR required: {}".format(file_type, requires_ocr),
                    level="info",
                )
            except Exception as e:
                self._db_rollback()
                errors.append("{}: {}".format(filename, e))
                self._log_activity(
                    filename, "Upload Failed",
                    str(e),
                    level="error",
                )

        lines = []
        if uploaded:
            lines.append("Uploaded ({}):\n  ".format(len(uploaded)) +
                         "\n  ".join(uploaded))
        if skipped:
            lines.append("Skipped ({}):\n  ".format(len(skipped)) +
                         "\n  ".join(skipped))
        if errors:
            lines.append("Errors ({}):\n  ".format(len(errors)) +
                         "\n  ".join(errors))

        if uploaded:
            messagebox.showinfo("Upload Complete", "\n\n".join(lines))
            self._refresh()
        else:
            messagebox.showwarning("Nothing Uploaded",
                                   "\n\n".join(lines) or "No files uploaded.")

    # =========================================================================
    #  BANK MANAGER
    # =========================================================================

    def _open_bank_manager(self):
        if not self._db:
            messagebox.showwarning("No Database",
                                   "Cannot manage banks: not connected to the database.")
            return
        # Reuse existing window if still open
        if self._bank_dialog and self._bank_dialog.winfo_exists():
            self._bank_dialog.lift()
            self._bank_dialog.focus_set()
            return
        self._bank_dialog = BankManagerDialog(self, self._db, self._db_rollback)

    # =========================================================================
    #  ACTIVITY LOG
    # =========================================================================

    def _init_activity_log(self):
        """Auto-create the activity_log table if it doesn't exist."""
        if not self._db:
            return
        try:
            self._db.cursor.execute("""
                CREATE TABLE IF NOT EXISTS activity_log (
                    log_id    SERIAL PRIMARY KEY,
                    logged_at TIMESTAMP NOT NULL DEFAULT NOW(),
                    invoice_ref VARCHAR(20),
                    action    VARCHAR(120) NOT NULL,
                    details   TEXT,
                    level     VARCHAR(20) DEFAULT 'info'
                )
            """)
            self._db.conn.commit()
        except Exception:
            try:
                self._db.conn.rollback()
            except Exception:
                pass

    def _log_activity(self, invoice_ref, action, details="", level="info"):
        """Append an entry to the in-memory log and persist to DB."""
        entry = {
            "time":        datetime.now(),
            "invoice_ref": invoice_ref or "",
            "action":      action,
            "details":     details,
            "level":       level,   # info | success | warning | error
        }
        self._activity_log.append(entry)

        if self._db:
            try:
                self._db.cursor.execute(
                    "INSERT INTO activity_log (invoice_ref, action, details, level) "
                    "VALUES (%s, %s, %s, %s)",
                    (invoice_ref, action, details, level)
                )
                self._db.conn.commit()
            except Exception:
                try:
                    self._db.conn.rollback()
                except Exception:
                    pass

        # Live-push to the log dialog if it's already open
        if self._log_dialog and self._log_dialog.winfo_exists():
            self._log_dialog.append_entry(entry)

    def _open_activity_log(self):
        if self._log_dialog and self._log_dialog.winfo_exists():
            self._log_dialog.lift()
            self._log_dialog.focus_set()
            return
        self._log_dialog = ActivityLogDialog(
            self, self._activity_log, self._db, self._db_rollback
        )

    # =========================================================================
    #  AI VALIDATION
    # =========================================================================

    def _run_ai_for_row(self, row):
        db_id = row.invoice.get("_db_id", row.invoice["invoice_id"])
        t = self._active_threads.get(db_id)
        if t and t.is_alive():
            return
        row.set_processing()
        self._log_activity(
            row.invoice.get("invoice_id", ""),
            "Validation Started",
            "AI agent dispatched for invoice {}".format(row.invoice.get("invoice_id", db_id)),
            level="info",
        )
        t = threading.Thread(target=self._validation_worker, args=(db_id, row), daemon=True)
        self._active_threads[db_id] = t
        t.start()

    def _run_selected(self):
        selected = [r for r in self._rows
                    if r.is_selected and r.invoice.get("validation_status") == "pending"]
        if not selected:
            messagebox.showinfo("Nothing Selected",
                                "Please select one or more pending invoices.")
            return
        for row in selected:
            self._run_ai_for_row(row)

    def _validation_worker(self, db_id, row):
        result = {"status": "error", "confidence": None,
                  "matched_ids": [], "message": "Agent not available",
                  "_agent_wrote_db": False}

        if self._agent:
            try:
                raw = self._agent.validate_transaction(db_id)

                if raw is None:
                    result = {"status": "error", "confidence": None,
                              "matched_ids": [], "message": "Agent returned None — check logs",
                              "_agent_wrote_db": False}

                elif isinstance(raw, dict) and raw.get("error"):
                    result = {"status": "error", "confidence": None,
                              "matched_ids": [],
                              "message": raw.get("message", "Unknown agent error"),
                              "_agent_wrote_db": False}

                elif isinstance(raw, dict) and "validated" in raw:
                    conf = raw.get("confidence")
                    txns = raw.get("transactions") or []
                    if raw["validated"]:
                        result = {
                            "status":          "auto_validated",
                            "confidence":      conf,
                            "matched_ids":     txns,
                            "_agent_wrote_db": True,
                            "message":         "",
                        }
                    else:
                        result = {
                            "status":          "needs_manual",
                            "confidence":      conf,
                            "matched_ids":     txns,
                            "_agent_wrote_db": False,
                            "message":         "",
                        }
                else:
                    result = {"status": "error", "confidence": None,
                              "matched_ids": [], "message": str(raw),
                              "_agent_wrote_db": False}

            except Exception as e:
                result = {"status": "error", "confidence": None,
                          "matched_ids": [], "message": str(e),
                          "_agent_wrote_db": False}
        else:
            # ── Simulation mode ── remove once real agent is connected ──────
            import time, random
            time.sleep(2.5)
            confidence = round(random.uniform(0.72, 0.99), 4)
            status = "auto_validated" if confidence >= CONFIDENCE_THRESHOLD else "needs_manual"
            result = {
                "status":          status,
                "confidence":      confidence,
                "matched_ids":     ["TXN-{}".format(random.randint(100000, 999999))],
                "_agent_wrote_db": False,
                "message":         "Simulated (agent not connected)",
            }

        self.after(0, self._on_validation_done, db_id, row, result)

    def _on_validation_done(self, db_id, row, result):
        status      = result.get("status", "error")
        confidence  = result.get("confidence")
        matched_ids = result.get("matched_ids") or []
        agent_wrote = result.get("_agent_wrote_db", False)

        row.set_result(status, confidence, matched_ids, result.get("message", ""))

        # ── Activity log ──────────────────────────────────────────────────────
        inv_ref = row.invoice.get("invoice_id", str(db_id))
        if status == "auto_validated":
            self._log_activity(
                inv_ref, "Auto-Validated",
                "Confidence {:.1f}% — matched transaction(s): {}".format(
                    (confidence or 0) * 100,
                    ", ".join(str(t) for t in matched_ids) or "none"),
                level="success",
            )
        elif status == "needs_manual":
            self._log_activity(
                inv_ref, "Needs Manual Review",
                "Confidence {:.1f}% — below auto-approval threshold.".format(
                    (confidence or 0) * 100),
                level="warning",
            )
            self._open_manual_review(row, auto_open=True)
        elif status == "error":
            self._log_activity(
                inv_ref, "Validation Error",
                result.get("message", "Unknown error"),
                level="error",
            )
        # ─────────────────────────────────────────────────────────────────────

        for inv in self._invoices:
            if inv.get("_db_id") == db_id:
                inv["validation_status"] = status
                inv["_confidence"]       = confidence
                inv["_matched_ids"]      = matched_ids
                break

        if self._db and status == "auto_validated" and confidence is not None and not agent_wrote:
            try:
                self._db.update_data(
                    "invoices", ["validation_status"], [True],
                    "invoice_id = %s", (db_id,)
                )
                self._db.insert_data(
                    "validation_details",
                    ["invoice_id", "confidence_score"],
                    [db_id, confidence]
                )
                for txn_id in matched_ids:
                    try:
                        self._db.insert_data(
                            "validation_transactions",
                            ["invoice_id", "transaction_id"],
                            [db_id, txn_id]
                        )
                    except Exception:
                        self._db_rollback()
            except Exception as e:
                self._db_rollback()
                print("[DB] Failed to persist auto_validated for {}: {}".format(db_id, e))

        self._active_threads.pop(db_id, None)
        self._update_stats()

    # =========================================================================
    #  MANUAL APPROVAL
    # =========================================================================

    def _approve_row(self, row):
        self._open_manual_review(row, auto_open=False)

    def _open_manual_review(self, row, auto_open=False):
        display_id  = row.invoice["invoice_id"]
        db_id       = row.invoice.get("_db_id", display_id)

        existing = self._manual_review_dialogs.get(db_id)
        if existing and existing.winfo_exists():
            existing.focus_set()
            existing.lift()
            return

        review_data = self._get_manual_review_data(row.invoice)
        dialog = ManualValidationDialog(
            self,
            row.invoice,
            review_data,
            on_approve=lambda r=row: self._commit_manual_approval(r),
        )
        dialog.protocol("WM_DELETE_WINDOW", dialog.destroy)
        self._manual_review_dialogs[db_id] = dialog
        dialog.bind("<Destroy>", lambda _event, key=db_id: self._manual_review_dialogs.pop(key, None))

        if auto_open:
            self._log_activity(
                display_id, "Manual Review Window Opened",
                "Database invoice and matched transaction data loaded for cross-check.",
                level="info",
            )

    def _get_manual_review_data(self, invoice):
        db_id = invoice.get("_db_id")
        data = {
            "invoice_row": None,
            "ocr_result": None,
            "transactions": [],
        }

        if not self._db or db_id is None:
            data["ocr_result"] = {
                "date": invoice.get("date_of_purchase"),
                "invoice_amount": invoice.get("amount"),
                "currency": invoice.get("currency"),
                "vendor": invoice.get("description"),
            }
            return data

        try:
            invoice_rows = self._db.retrieve_data(
                "invoices",
                condition="invoice_id = %s",
                condition_values=(db_id,)
            )
            if not invoice_rows:
                invoice_rows = self._db.retrieve_data(
                    "invoices",
                    condition="invoice_ID = %s",
                    condition_values=(db_id,)
                )
            data["invoice_row"] = invoice_rows[0] if invoice_rows else None
        except Exception:
            self._db_rollback()

        try:
            ocr_rows = self._db.retrieve_data(
                "ocr_results",
                condition="invoice_id = %s",
                condition_values=(db_id,)
            )
            if not ocr_rows:
                ocr_rows = self._db.retrieve_data(
                    "ocr_results",
                    condition="invoice_ID = %s",
                    condition_values=(db_id,)
                )
            if ocr_rows:
                import json as _json
                result = ocr_rows[0].get("ocr_result") or ocr_rows[0].get("OCR_result")
                if isinstance(result, str):
                    result = _json.loads(result)
                data["ocr_result"] = result
        except Exception:
            self._db_rollback()

        matched_ids = invoice.get("_matched_ids") or []
        if matched_ids:
            try:
                data["transactions"] = self._db.retrieve_data(
                    "transactions",
                    condition="transaction_id IN %s",
                    condition_values=(tuple(matched_ids),)
                )
            except Exception:
                self._db_rollback()
                try:
                    data["transactions"] = self._db.retrieve_data(
                        "transactions",
                        condition="transaction_ID IN %s",
                        condition_values=(tuple(matched_ids),)
                    )
                except Exception:
                    self._db_rollback()

        return data

    def _commit_manual_approval(self, row):
        display_id  = row.invoice["invoice_id"]
        db_id       = row.invoice.get("_db_id", display_id)
        confidence  = row.invoice.get("_confidence") or 0.0
        matched_ids = row.invoice.get("_matched_ids") or []

        row.set_approved()
        for inv in self._invoices:
            if inv.get("_db_id") == db_id:
                inv["validation_status"] = "manually_validated"
                break

        if self._db:
            try:
                self._db.update_data(
                    "invoices", ["validation_status"], [True],
                    "invoice_id = %s", (db_id,)
                )
                self._db.insert_data(
                    "validation_details",
                    ["invoice_id", "confidence_score"],
                    [db_id, confidence]
                )
                for txn_id in matched_ids:
                    try:
                        self._db.insert_data(
                            "validation_transactions",
                            ["invoice_id", "transaction_id"],
                            [db_id, txn_id]
                        )
                    except Exception:
                        self._db_rollback()
                self._log_activity(
                    display_id, "Manually Approved",
                    "Confidence {:.1f}% | Matched: {}".format(
                        confidence * 100,
                        ", ".join(str(t) for t in matched_ids) or "none"),
                    level="success",
                )
            except Exception as e:
                self._db_rollback()
                messagebox.showerror("DB Error", "Could not update database:\n{}".format(e))
                self._log_activity(display_id, "Approval DB Error", str(e), level="error")
        else:
            self._log_activity(display_id, "Manually Approved", "No DB — in-memory only",
                               level="success")

        self._update_stats()

    # =========================================================================
    #  TOOLBAR INTERACTIONS
    # =========================================================================

    def _on_search(self, _event=None):
        self._render_invoices()

    def _on_filter_change(self, _value=None):
        self._render_invoices()

    def _refresh(self):
        self._load_invoices()

    def _on_select_all(self):
        val = self._select_all_var.get()
        for row in self._rows:
            row.set_selected(val)

    def _update_select_all_state(self):
        if self._rows:
            self._select_all_var.set(all(r.is_selected for r in self._rows))

    # =========================================================================
    #  CLEANUP
    # =========================================================================

    def on_close(self):
        if self._db:
            try:
                self._db.close_connection()
            except Exception:
                pass
        self.destroy()


# =============================================================================
#  ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    app = GlobalTreasuryApp()
    app.protocol("WM_DELETE_WINDOW", app.on_close)
    app.mainloop()
